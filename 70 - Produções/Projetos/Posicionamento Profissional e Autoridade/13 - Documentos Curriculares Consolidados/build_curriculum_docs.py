from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path("deliverables")
OUT.mkdir(exist_ok=True)

INK = RGBColor(26, 45, 63)
BLUE = RGBColor(43, 91, 126)
TEAL = RGBColor(42, 120, 124)
MUTED = RGBColor(92, 104, 113)
LIGHT = "EAF1F4"
VERY_LIGHT = "F5F7F8"
WHITE = "FFFFFF"


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_fixed_table_layout(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    total_twips = sum(int(round(width * 1440)) for width in widths)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    # Compensa o recuo visual produzido pela margem interna inicial da célula,
    # mantendo a borda externa alinhada ao corpo do texto em diferentes versões do Word.
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width * 1440))))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
            tc_pr = row.cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(round(width * 1440))))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_doc(doc, preset="narrative"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.88)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7 if preset == "narrative" else 5)
    normal.paragraph_format.line_spacing = 1.24 if preset == "narrative" else 1.15

    for style_name, size, before, after, color in [
        ("Heading 1", 16, 18, 8, BLUE),
        ("Heading 2", 13, 13, 6, BLUE),
        ("Heading 3", 11.5, 9, 4, TEAL),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        st = styles[list_style]
        st.font.name = "Calibri"
        st.font.size = Pt(10.7)
        st.paragraph_format.left_indent = Inches(0.38)
        st.paragraph_format.first_line_indent = Inches(-0.18)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15

    if "Lead" not in styles:
        lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Lead"]
    lead.font.name = "Calibri"
    lead.font.size = Pt(12.2)
    lead.font.color.rgb = INK
    lead.paragraph_format.space_after = Pt(11)
    lead.paragraph_format.line_spacing = 1.28

    if "Small Note" not in styles:
        sn = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        sn = styles["Small Note"]
    sn.font.name = "Calibri"
    sn.font.size = Pt(9.2)
    sn.font.color.rgb = MUTED
    sn.paragraph_format.space_after = Pt(4)
    sn.paragraph_format.line_spacing = 1.1

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("FABIANO DELIBERALLI  |  CURRÍCULO 2026")
    set_font(hr, size=8.5, bold=True, color=MUTED)
    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title_block(doc, title, subtitle, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    set_font(r, size=28, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(13)
    r = p.add_run(subtitle)
    set_font(r, size=15, bold=False, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(label)
    set_font(r, size=9.5, bold=True, color=MUTED)


def add_signature(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Corpo, trauma, consciência e espiritualidade na compreensão e tradução da experiência humana.")
    set_font(r, size=12.5, bold=True, italic=True, color=TEAL)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            label, body = item
            r = p.add_run(label)
            set_font(r, size=10.7, bold=True, color=INK)
            r = p.add_run(body)
            set_font(r, size=10.7, color=INK)
        else:
            r = p.add_run(item)
            set_font(r, size=10.7, color=INK)


def add_key_value_table(doc, rows, widths=(1.55, 4.95)):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    header_row = table.rows[0]
    prevent_row_split(header_row)
    set_repeat_table_header(header_row)
    for index, heading in enumerate(("Campo", "Informação")):
        cell = header_row.cells[index]
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(heading)
        set_font(run, size=9.4, bold=True, color=INK)
    for label, value in rows:
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        set_cell_shading(cells[0], LIGHT)
        for c in cells:
            set_cell_margins(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=9.4, bold=True, color=INK)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=9.4, color=INK)
    set_fixed_table_layout(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_records_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell, top=95, bottom=95)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=8.5, bold=True, color=INK)
    for row_idx, row in enumerate(rows):
        data_row = table.add_row()
        prevent_row_split(data_row)
        cells = data_row.cells
        if row_idx % 2:
            for cell in cells:
                set_cell_shading(cell, VERY_LIGHT)
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(value)
            set_font(r, size=8.4, color=INK)
    set_fixed_table_layout(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def build_curriculum():
    doc = Document()
    configure_doc(doc, "narrative")
    add_title_block(
        doc,
        "Fabiano Deliberalli",
        "Psicólogo clínico e psicoterapeuta  |  CRP 06/98630",
        "CURRÍCULO PROFISSIONAL AMPLIADO · VERSÃO CONSOLIDADA PARA REVISÃO · 21/08/2026",
    )
    add_signature(doc)

    doc.add_heading("Perfil profissional", level=1)
    p = doc.add_paragraph(style="Lead")
    p.add_run(
        "Fabiano Deliberalli é psicólogo clínico e psicoterapeuta, com mais de 30 anos de trajetória no cuidado terapêutico, iniciada nas práticas corporais e integrativas e consolidada com a Psicologia."
    )
    doc.add_paragraph(
        "Seu percurso articula corpo, escuta psicanalítica, Psicologia, trauma, neurociência, consciência e espiritualidade. A diversidade das formações é organizada por uma questão central: como compreender o ser humano em suas dimensões corporal, afetiva, relacional, simbólica, consciente e espiritual sem reduzi-lo a uma única linguagem?"
    )
    doc.add_paragraph(
        "Na clínica, dedica-se à compreensão de padrões, sofrimento, trauma, processos de regulação e experiências espirituais, existenciais ou incomuns que pedem escuta cuidadosa. Seu trabalho procura traduzir a experiência vivida e favorecer integração, preservando rigor clínico, abertura e discernimento."
    )

    doc.add_heading("Eixos de atuação", level=1)
    add_bullets(doc, [
        "Psicologia clínica e psicoterapia exclusivamente on-line.",
        "Trauma, reprocessamento, regulação e integração da experiência.",
        "Corpo, percepção, respiração, movimento e experiência incorporada.",
        "Consciência, espiritualidade, sentido e experiências humanas complexas.",
        "Educação continuada, palestras e tradução pública de temas clínicos e interdisciplinares.",
    ])

    doc.add_heading("Formação acadêmica", level=1)
    add_bullets(doc, [
        ("Psicologia — Universidade São Marcos. ", "Curso concluído em 4 de fevereiro de 2010; diploma emitido em 14 de setembro de 2010."),
        ("Pós-graduação em Psicologia Transpessoal — FACIS, Faculdade de Ciências da Saúde de São Paulo. ", "Formação concluída em 2010, com Trabalho de Conclusão de Curso entregue e certificado emitido em 2011. Monografia orientada pela Profª Drª Márcia Tabone."),
        ("Pós-graduação lato sensu em Neurociências e Comportamento — PUCRS. ", "Concluída em 2026; certificado emitido em 25 de março de 2026."),
        ("TCC da PUCRS. ", "Sincronia Neural Interpessoal e Estados Ampliados de Consciência: uma abordagem neurobiológico-relacional da espiritualidade, aprovado com nota 9,5."),
    ])
    doc.add_paragraph(
        "Interesses atuais de estudo: neurofenomenologia, experiência incorporada, trauma, espiritualidade, sincronia interpessoal, consciência e diferenciação clínica de experiências incomuns."
    )

    doc.add_heading("Formação clínica e psicoterapêutica", level=1)
    add_bullets(doc, [
        ("Psicanálise Integrativa. ", "Duas formações: uma na SBPI/Psicanálise Integrativa Clínica Escola, com Maria de Fátima Mora, e outra no Instituto Brasileiro de Transpsicanálise, com André Keppe como professor principal. A primeira possui certificado de 23/07/2002 e declaração curricular de 31/07/2002, com 294 horas; os dois documentos comprovam a mesma formação. A formação psicanalítica integrou estágios e o início da prática psicoterapêutica."),
        ("EMDR. ", "Formação continuada desde 2010, com treinamento básico, cursos clínicos, supervisão e atualizações em trauma, dissociação, estados do ego, apego e reprocessamento. Inclui cursos com Sandra Paulsen sobre trauma precoce e apego (7h, 2015) e EMDR e Estados do Ego (14h, 2019)."),
        ("Brainspotting. ", "Formação continuada com Fases 1 a 5 concluídas entre 2011 e 2025. A Fase 5 foi um treinamento avançado de 24 horas com David Grand; a Fase 3 foi concluída na sequência formativa entre as Fases 2 e 4."),
        ("Meditação Bodyspotting. ", "Workshop complementar de 4 horas com Marília Toscano, realizado em 2026 e certificado pela ABBSP."),
        ("Psicoterapia de Integração e Reprocessamento do Trauma — Instituto Aleces. ", "Programa de 216 horas, realizado entre 2020 e 2023, com Mário Salvador e Carmen Cuenca."),
        ("Supervisão clínica continuada. ", "Supervisão mensal em grupo com Mário Salvador, em continuidade em 2026; formação presencial documentada em Brasília em 2018, atividade formativa em 2019 e imersões presenciais de três dias em 2025 e 2026."),
        ("Neurobiologia Interpessoal. ", "Curso de 16 horas com Daniel Siegel, em 2021, além de atualizações anteriores documentadas."),
        ("Teoria Polivagal aplicada à prática clínica. ", "Curso de 12 horas com Deb Dana, concluído em setembro de 2023."),
        ("Hipnose clínica e terapêutica. ", "Formação continuada com estudos junto a Alberto Dell’Isola, Marco Ceda Natali, Fernando Toledo Piza e Sofia Bauer. Registros incluem Mirroring Hands (12h, 2019), Professional & Clinical Hypnosis B1–B3 (30h, 2021), B4–B8 (70h, 2022) e Hypnotic Mastery Bootcamp (18h, 2022), além de curso on-line com Sofia Bauer, com acesso pago confirmado entre agosto de 2023 e fevereiro de 2024."),
    ])

    doc.add_heading("Trajetória corporal e integrativa", level=1)
    doc.add_paragraph(
        "Antes da graduação em Psicologia, Fabiano construiu experiência extensa com práticas corporais, reabilitação, massagem, Medicina Tradicional Chinesa, acupuntura, Qi Gong, respiração, movimento e terapias integrativas. Essa etapa constitui a base incorporada a partir da qual passou a estudar escuta, sofrimento, consciência e mudança."
    )
    add_bullets(doc, [
        ("Técnico em Reabilitação — modalidade Massagista — EOMA. ", "Formação concluída em 1995, com 1.440 horas, incluindo 162 horas de estágio."),
        ("Shiozawa Prevenção em Saúde. ", "Curso Intensivo de Shiatsu de 100 horas, realizado de junho a dezembro de 1994, e estágio documentado de 2.000 horas, realizado de janeiro de 1995 a julho de 1996. Curso e estágio são registros separados. A experiência prosseguiu em atendimento e estudos corporais durante aproximadamente dez anos."),
        ("Acupuntura Tradicional Chinesa — CEMETRAC. ", "Formação de 1.200 horas, realizada entre 2002 e 2004, sob direção de Mestre Liu Chih Ming."),
        ("Medicina Tradicional Chinesa. ", "Estudos com Mestre Liu Chih Ming, Ernesto Garcia e Pedro Pablo Arias Capdet, pela Pró-Salus Vitae, além de programas nacionais e internacionais de formação continuada."),
        ("Máster Iberoamericano en Acupuntura Bioenergética y Moxibustión. ", "Concluído e certificado em outubro de 2015 no âmbito institucional da Universidad de Medicina Tradicional China de Yunnan, do CEMETC e da Fundación Europea de Medicina Tradicional China."),
        ("CEMETC. ", "Formação internacional continuada em Acupuntura Bioenergética y Moxibustión, com níveis 1, 2 e 3, programas curriculares documentados e certificados gerais de 360 horas/14,4 créditos ECTS."),
        ("Venia Docente do CEMETC. ", "Credencial institucional recebida em Valladolid, em 8 de novembro de 2014, para ministrar programas de formação em Acupuntura Bioenergética estabelecidos pela instituição — marco docente de sua trajetória em MTC."),
        ("Qi Gong, respiração e movimento. ", "Formações em práticas de regulação, presença corporal, atenção e movimento, incluindo formação de instrutor e práticas específicas."),
        ("Associação Palas Athena. ", "Sequência completa de três módulos de Atenção e concentração nas práticas meditativas; módulos 2 e 3 documentados em 2013."),
        ("Seitai, Quiropraxia e práticas corporais orientais. ", "Estudos, prática cotidiana e experiência didática integrados à trajetória corporal iniciada no Shiozawa."),
        ("Bioenergologia e Bioenergopatia. ", "Curso livre concluído em agosto de 2004, com nota final 9,0 e programa curricular de 114 unidades-aula documentadas, sem equivalência horária total declarada. Integra a genealogia de estudos sobre consciência, percepção e práticas energéticas."),
        ("Reiki e práticas energéticas. ", "Formação, prática e ensino em período anterior à docência em Psicologia."),
    ])

    doc.add_heading("Comunicação, linguagem e processos de mudança", level=1)
    add_bullets(doc, [
        ("Programação Neurolinguística — SBPNL. ", "Formação Practitioner realizada antes das formações posteriores no Trivium."),
        ("Trivium. ", "Practitioner em PNL de 112 horas e Master Practitioner de 80 horas, ambos documentados em 2001."),
    ])
    doc.add_paragraph(
        "A PNL integra a genealogia de estudos sobre comunicação, linguagem, percepção de padrões e processos de mudança, em diálogo com a trajetória posterior em Psicologia, psicoterapia e trauma."
    )

    doc.add_heading("Experiência clínica e institucional", level=1)
    doc.add_paragraph(
        "Fundou o Bioquantum, posteriormente Centralma, onde realizou atendimentos, desenvolveu práticas corporais e recebeu profissionais para cursos e atividades formativas. As atividades presenciais foram encerradas em 2020. A Centralma permanece como empresa atual e como eixo administrativo dos atendimentos psicológicos exclusivamente on-line."
    )
    doc.add_paragraph(
        "Sua atuação clínica atual se organiza em torno de Psicologia, psicoterapia, trauma, regulação, consciência, espiritualidade, experiências existenciais e experiências incomuns. Os recursos são selecionados e integrados conforme a experiência, a história e as necessidades da pessoa atendida."
    )

    doc.add_heading("Docência, palestras e educação continuada", level=1)
    add_bullets(doc, [
        "Cursos de Reiki e formação de alunos em práticas integrativas.",
        "Curso sobre Seitai e Quiropraxia durante a atuação no Shiozawa.",
        "Aula convidada na Faculdade de Psicologia da PUC-SP, em 1º de outubro de 2002: Corpo e Mente: um mesmo sistema.",
        "Professor assistente de Gilda Moura em turma presencial realizada no Centralma, em 2018/2019.",
        "Ministrante, em parceria com Gilda Moura, de três cursos on-line de auto-hipnose entre 2020 e 2025.",
        "Professor e cocriador de Além das Fronteiras do Saber, com cinco edições on-line a partir de 2020.",
        "Palestras e cursos em organizações, eventos e projetos de desenvolvimento humano.",
        "Workshop da ABBSP, em 7 de março de 2026, sobre autorregulação, presença clínica, processos espirituais e estados ampliados de consciência.",
        "Palestra no encontro de experienciadores do IGM, em agosto de 2026, com Daiana Leite.",
        "Professor do eixo Consciência no CEEP — Comunidade de Estudos em Espiritualidade e Psicoterapia.",
        "Criador e professor do projeto Traduzindo o Ser Humano.",
    ])
    p = doc.add_paragraph()
    r = p.add_run("Formulação docente: ")
    set_font(r, bold=True, color=INK)
    p.add_run("professor, palestrante e facilitador em cursos livres e educação continuada.")

    doc.add_heading("Campo autoral e projetos atuais", level=1)
    doc.add_heading("CEEP — Comunidade de Estudos em Espiritualidade e Psicoterapia", level=2)
    doc.add_paragraph(
        "Projeto de educação continuada para profissionais de saúde, organizado nos eixos Clínica, Ciência e Consciência. Fabiano contribui principalmente no eixo Consciência, com fenomenologia, neurofenomenologia, significado, discernimento, experiências espirituais, estados ampliados de consciência e diferenciação entre experiência vivida, interpretação, hipótese, evidência e conclusão."
    )
    doc.add_heading("Traduzindo o Ser Humano", level=2)
    doc.add_paragraph(
        "Projeto educacional e de produtos digitais dedicado a tornar compreensíveis temas complexos sobre corpo, trauma, consciência, espiritualidade, padrões e transformação, preservando rigor e linguagem acessível."
    )
    doc.add_heading("Produção acadêmica e editorial", level=2)
    doc.add_paragraph(
        "O TCC da PUCRS constitui a base de uma produção autoral em desenvolvimento sobre sincronia neural interpessoal, espiritualidade, consciência e uma abordagem neurobiológico-relacional da experiência. Entre as próximas etapas estão a preparação de artigo, a organização do Lattes, do ORCID e de uma bibliografia pública."
    )

    doc.add_heading("Perspectiva de trabalho", level=1)
    p = doc.add_paragraph(style="Lead")
    r = p.add_run(
        "O corpo e as práticas integrativas formaram a base incorporada; a Psicanálise organizou a escuta; a Psicologia formalizou a identidade profissional; o trauma e a neurobiologia aprofundaram a clínica; a fenomenologia e a consciência deram direção autoral; a docência transforma esse percurso em contribuição pública."
    )
    r.italic = True
    doc.add_paragraph(
        "Sua postura profissional é orientada pela possibilidade de acolher a experiência sem patologizá-la automaticamente, investigar sem confirmar previamente uma interpretação e traduzir fenômenos complexos sem reduzi-los a uma única explicação."
    )

    doc.add_heading("Mini bio consolidada", level=1)
    doc.add_paragraph(
        "Fabiano Deliberalli é psicólogo clínico e psicoterapeuta (CRP 06/98630), com mais de 30 anos de trajetória no cuidado terapêutico, iniciada nas práticas corporais e integrativas e consolidada com a Psicologia. É pós-graduado em Psicologia Transpessoal e em Neurociências e Comportamento pela PUCRS, com formações em Psicanálise Integrativa, EMDR, Brainspotting, hipnose clínica e terapêutica, Neurobiologia Interpessoal e trauma."
    )
    doc.add_paragraph(
        "Seu trabalho articula corpo, trauma, consciência e espiritualidade para compreender, traduzir e favorecer a integração de experiências humanas complexas, com rigor clínico, abertura e discernimento. É professor e palestrante em cursos livres e educação continuada e desenvolve o CEEP e o projeto Traduzindo o Ser Humano."
    )

    doc.add_heading("Nota de uso", level=1)
    doc.add_paragraph(
        "Esta é a versão pública ampliada. Dela devem ser derivadas versões específicas para clínica, CEEP, academia, mídia, palestras, site e produtos digitais, preservando identidade, formulação temporal e assinatura de posicionamento, com seleção de ênfases conforme o contexto."
    )
    p = doc.add_paragraph(style="Small Note")
    p.add_run("Fonte de consolidação: vault Obsidian sincronizado no GitHub, branch de revisão curricular de 21/08/2026.")

    path = OUT / "Curriculo_Profissional_Consolidado_Fabiano_Deliberalli_2026.docx"
    doc.save(path)
    return path


def build_dossier():
    doc = Document()
    configure_doc(doc, "compact")
    section = doc.sections[0]
    section.header.paragraphs[0].clear()
    hr = section.header.paragraphs[0].add_run("FABIANO DELIBERALLI  |  DOSSIÊ CURRICULAR MESTRE")
    set_font(hr, size=8.5, bold=True, color=MUTED)
    add_title_block(
        doc,
        "Dossiê Curricular Mestre",
        "Fabiano Deliberalli  |  fonte interna para currículo, site, CEEP, mídia e projetos",
        "VERSÃO CONSOLIDADA PARA REVISÃO · 21/08/2026",
    )
    add_signature(doc)

    doc.add_heading("1. Identidade profissional e regra de leitura", level=1)
    add_key_value_table(doc, [
        ("Nome profissional", "Fabiano Deliberalli"),
        ("Identidade atual", "Psicólogo clínico e psicoterapeuta"),
        ("Registro", "CRP 06/98630"),
        ("Linha temporal", "Mais de 30 anos de trajetória no cuidado terapêutico, iniciada nas práticas corporais e integrativas e consolidada com a Psicologia."),
        ("Posicionamento", "Corpo, trauma, consciência e espiritualidade na compreensão e tradução da experiência humana."),
        ("Núcleo autoral", "Traduzir experiências humanas complexas sem reduzi-las prematuramente a diagnóstico, crença ou explicação única."),
    ])
    doc.add_paragraph(
        "Este dossiê preserva a totalidade relevante da trajetória e distingue formação, supervisão, experiência clínica, docência, vínculo institucional, produção acadêmica e credenciais históricas. Confirmações diretas de Fabiano integram a fonte curricular; a ausência momentânea de arquivo digital não cria pendência de formação."
    )

    doc.add_heading("2. Tese de coerência da trajetória", level=1)
    doc.add_paragraph(
        "A amplitude do percurso não deve ser apresentada como uma coleção de técnicas equivalentes. Sua coerência está em uma pergunta que atravessa décadas: como compreender o ser humano em sua dimensão corporal, afetiva, relacional, simbólica, consciente e espiritual sem reduzi-lo a uma única linguagem?"
    )
    add_bullets(doc, [
        ("Corpo. ", "Massoterapia, reabilitação, práticas orientais, MTC, Qi Gong e respiração."),
        ("Escuta e linguagem. ", "Psicanálise Integrativa, psicoterapia breve, Winnicott, hipnose e PNL."),
        ("Psicologia e integração. ", "Graduação em Psicologia e pós-graduação em Psicologia Transpessoal."),
        ("Trauma e regulação. ", "EMDR, Brainspotting, Aleces, Neurobiologia Interpessoal, Teoria Polivagal e supervisão continuada."),
        ("Consciência e autoria. ", "Neurociências e Comportamento, TCC sobre sincronia neural e estados ampliados de consciência, CEEP, palestras e Traduzindo o Ser Humano."),
    ])

    doc.add_heading("3. Eixo 1 — Psicologia, Psicanálise e formação acadêmica", level=1)
    rows = [
        ("Graduação em Psicologia", "Universidade São Marcos", "Conclusão 04/02/2010; diploma 14/09/2010", "—", "Formação acadêmica principal; CRP desde 2010"),
        ("Trabalho final da graduação", "Universidade São Marcos", "Defesa 04/12/2009", "—", "O normal e o patológico nas vivências espirituais: explorações psicanalíticas e transpessoais"),
        ("Pós-graduação em Psicologia Transpessoal", "FACIS; orientação Profª Drª Márcia Tabone", "Conclusão 2010; certificado 2011", "—", "TCC entregue; formação concluída"),
        ("Pós-graduação em Neurociências e Comportamento", "PUCRS", "Certificado 25/03/2026", "—", "TCC aprovado com nota 9,5"),
        ("Psicanálise Integrativa — 1ª formação", "SBPI / Clínica Escola; Maria de Fátima Mora", "Certificado 23/07/2002; declaração 31/07/2002", "294h", "Documentos complementares; estágios e início da prática psicoterapêutica"),
        ("Psicanálise Integrativa — 2ª formação", "Instituto Brasileiro de Transpsicanálise; André Keppe", "Período inicial da trajetória clínica", "—", "Formação confirmada pelo autor e apoiada por registros complementares"),
        ("Neurociência das Funções Mentais", "Transpsicanálise", "02/09/2002", "12h", "Curso complementar"),
        ("Bases Bioquímicas dos Transtornos Mentais", "Transpsicanálise", "16/09–16/12/2002", "36h", "Curso complementar"),
    ]
    add_records_table(doc, ["Formação", "Instituição / professor", "Data / período", "Carga", "Natureza e uso"], rows, [1.45, 1.35, 1.22, 0.55, 1.83])

    doc.add_heading("4. Eixo 2 — Trauma, reprocessamento, neurobiologia e regulação", level=1)
    rows = [
        ("EMDR — Básico nível I", "Formação EMDR", "07/05/2010", "20h", "Registro formativo"),
        ("EMDR — Intermediário", "Formação EMDR", "24/10/2010", "20h", "Registro formativo"),
        ("Treinamento Básico em EMDR", "Formação EMDR", "18–20/03/2011", "50h + 10h supervisão", "Treinamento teórico-prático e supervisão"),
        ("EMDR crianças, adolescentes, trauma e dissociação", "Formação EMDR", "18/09/2011", "28h", "Aprofundamento clínico"),
        ("Teoria e Psicoterapia dos Estados do Ego", "Brasília", "12/08/2012", "18h + 18 créditos", "Aprofundamento clínico"),
        ("EMDR Clínico e supervisão de casos", "Brasília", "03/02/2013", "20h + 20 créditos; 20h supervisão", "Manejo teórico-prático e supervisão"),
        ("Quando não há Palavras — trauma precoce e apego", "Sandra Paulsen", "02/11/2015", "7h", "III Congresso Brasileiro de EMDR"),
        ("Protocolos para atenção precoce e trauma continuado", "Formação EMDR", "16/02/2019", "14h", "Atualização"),
        ("EMDR e Estados do Ego", "Sandra Paulsen / TraumaClinic", "04–05/04/2019", "14h", "Atualização clínica"),
        ("Atualizações em Neurobiologia Interpessoal", "São Paulo", "10/11/2013", "16h + 16 créditos", "Atualização relacionada a trauma e vínculo"),
        ("Brainspotting — Fases 1 e 2", "Formação Brainspotting", "03/04/2011", "14h + 14h", "Início da sequência formativa"),
        ("Brainspotting — Fases 1 e 2", "Formação Brainspotting", "10–15/08/2017", "—", "Nova sequência documentada"),
        ("Brainspotting — Fase 3", "Formação Brainspotting", "Entre Fases 2 e 4", "—", "Concluída e confirmada diretamente por Fabiano"),
        ("Brainspotting — Fase 4", "Formação Brainspotting", "02–04/08/2019", "18h", "Formação avançada"),
        ("Brainspotting — Fase 5", "David Grand / ABBSP", "22–24/08/2025", "24h", "Treinamento avançado on-line ao vivo"),
        ("Meditação Bodyspotting", "Marília Toscano / ABBSP", "11/04/2026", "4h", "Workshop complementar; não é uma fase"),
        ("Psicoterapia Integradora e Reprocessamento do Trauma", "Mário Salvador / AIBAPT", "05–08/09/2019", "28h", "Formação presencial"),
        ("Instituto Aleces — módulo 2", "Mário Salvador", "25–28/10/2018", "28h", "Formação presencial em Brasília"),
        ("Programa Aleces — Níveis I, II e III", "Mário Salvador e Carmen Cuenca", "2020–2023", "216h", "Programa completo; Modelo Aleceia"),
        ("Programa Aleces — Nível II", "Instituto Aleces", "2021–2022", "90h", "Componente incluído nas 216h"),
        ("Supervisão mensal em grupo", "Mário Salvador", "Continuidade em 2026", "—", "Formação clínica continuada"),
        ("Imersões presenciais", "Mário Salvador", "2025 e 2026", "3 dias cada", "Atendimentos e supervisões em período integral"),
        ("Neurobiologia Interpessoal", "Daniel Siegel", "16–17/10/2021", "16h", "Webconferência certificada"),
        ("Teoria Polivagal aplicada à prática clínica", "Deb Dana", "Setembro/2023", "12h", "Navigating the Nervous System"),
        ("Play of Life — visualizando mudanças", "Formação complementar", "28–29/04/2016", "12h", "Recurso clínico"),
        ("Play of Life — técnicas avançadas", "Formação complementar", "30/04–01/05/2016", "12h", "Grupos e famílias"),
        ("Neuroendocrinologia do Trauma", "Brasília", "14/11/2010", "14h", "Aprofundamento"),
        ("Transmissão Transgeracional de Trauma", "São Paulo", "18–19/03/2016", "16h", "Aprofundamento"),
    ]
    add_records_table(doc, ["Formação", "Instituição / professor", "Data / período", "Carga", "Natureza e uso"], rows, [1.45, 1.35, 1.22, 0.70, 1.68])

    doc.add_heading("5. Eixo 3 — Corpo, Reabilitação, MTC, acupuntura, Qi Gong e movimento", level=1)
    rows = [
        ("Técnico em Reabilitação — Massagista", "EOMA", "Conclusão 1995; emissão 14/07/1995", "1.440h; estágio 162h", "Marco inicial do percurso corporal"),
        ("Manipulação de clavícula e glenoumeral", "EOMA", "26/11/1994", "8h", "Aperfeiçoamento corporal"),
        ("Manipulação da bacia", "EOMA", "15/10/1994", "8h", "Aperfeiçoamento corporal"),
        ("Nós de Fitoterapia", "EOMA", "10/05/1994", "6h", "Curso complementar"),
        ("Alinhamento Postural, Reflexologia e Relaxamento", "EOMA", "26/03–25/06/1995", "60h", "Curso corporal"),
        ("Reflexologia e Alinhamento da Coluna", "EOMA", "Mar–jul/1996", "60h", "Curso corporal"),
        ("Drenagem Linfática", "EOMA", "Abr–jun/1996", "20h", "Curso corporal"),
        ("Curso Intensivo de Shiatsu", "Shiozawa Prevenção em Saúde", "Jun–dez/1994", "100h", "Formação corporal documentada"),
        ("Estágio em Shiatsu / práticas corporais", "Shiozawa Prevenção em Saúde", "Jan/1995–jul/1996", "2.000h", "Estágio documentado; separado do curso"),
        ("Atuação corporal e estudos", "Shiozawa / Mariko Sato", "Aprox. dez anos", "—", "Continuidade autobiográfica; Seitai e Quiropraxia"),
        ("Sotai / Reeducação Postural Integral", "Shiatsu Yasuragi / Japan Sotai Institute", "Fev/2015", "—", "Nível básico"),
        ("Chi Kung e Massagem Quiroprática/Seitai", "EOMA / Instituto Shioda", "01–07/07/2017", "—", "Com energização"),
        ("Massagem Quiroprática/Seitai — Nível II", "Mestre Kenichi Shioda", "10–16/02/2018", "—", "Aprofundamento"),
        ("Acupuntura Tradicional Chinesa", "CEMETRAC / Mestre Liu Chih Ming", "16/03/2002–29/02/2004", "1.200h", "Formação extensa em MTC"),
        ("Semiologia e Propedêutica da MTC", "Pró-Salus / Ernesto Garcia", "12/02–04/12/2005", "143h", "Aprofundamento"),
        ("Auriculoterapia", "Pró-Salus", "18/02–02/04/2006", "48h", "Curso básico"),
        ("Psiconeuroacupuntura Básica", "Instituto Juan Pablo Moltó", "07/07/2014", "18h", "On-line, Alicante"),
        ("Máster Iberoamericano en Acupuntura Bioenergética y Moxibustión", "Yunnan / CEMETC / Fundación Europea de MTC", "Certificado out/2015", "—", "Título conclusivo de formação continuada internacional"),
        ("CEMETC — níveis 1, 2 e 3", "Carlos Nogueira Pérez", "2014–2016", "360h / 14,4 ECTS", "Acupuntura Bioenergética y Moxibustión"),
        ("Venia Docente", "CEMETC", "Valladolid, 08/11/2014", "—", "Marco docente institucional para programas específicos"),
        ("Módulo internacional CEMETC", "Heiwa / CEMETC", "29–30/03/2014", "16h", "Formação continuada"),
        ("Módulos Long Tao / CEMETC", "CEMETC", "2015–2016", "32h cada", "Módulos I, II e III documentados"),
        ("Seminários clínicos CEMETC", "CEMETC", "2014", "6h a 11h", "Metabolismo, semiologia, dor, biomedições, oncologia, ginecologia e ambulatório"),
        ("Formação de instrutor de Qi Gong", "Programas documentados", "—", "96h", "Ba Duan Jin, Yi Jin Jing, Wu Qin Xi e Liu Zi Jue"),
        ("Chinese Health Qigong", "CHQA", "Data parcial", "1 Duan", "Certificado técnico nº 20140187"),
        ("Liu Zi Jue / Seis Sons de Cura", "Miguel Martín Vizán", "23–24/08/2014", "—", "Formação técnica"),
        ("Tai Chi Qigong Instructor Training — nível I", "Formação internacional", "Jan/2016", "—", "Formação de instrutor"),
        ("Sistema Energético Humano", "Programa em módulos", "Conclusão 07/12/2014", "48h", "Formação complementar"),
    ]
    add_records_table(doc, ["Formação", "Instituição / professor", "Data / período", "Carga", "Natureza e uso"], rows, [1.45, 1.35, 1.22, 0.72, 1.66])

    doc.add_heading("6. Eixo 4 — Hipnose, PNL, comunicação, linguagem e mudança", level=1)
    rows = [
        ("Hipnose clínica e terapêutica", "Alberto Dell’Isola, Marco Ceda Natali, Fernando Toledo Piza e Sofia Bauer", "Formação continuada", "—", "Eixo explícito de psicoterapia, linguagem e mudança"),
        ("Mirroring Hands — Essential Practice & Practitioner", "Richard Hill", "07–08/12/2019", "12h", "Registro documental"),
        ("Professional & Clinical Hypnosis B1–B3", "Formação em hipnose", "10–12/09/2021", "30h", "Registro documental"),
        ("Professional & Clinical Hypnosis B4–B8", "Formação em hipnose", "09–15/07/2022", "70h", "Registro documental"),
        ("Hypnotic Mastery Bootcamp", "Formação em hipnose", "06–07/08/2022", "18h", "Registro documental"),
        ("Curso on-line de hipnose", "Sofia Bauer", "Acesso ago/2023–fev/2024", "—", "Acesso pago confirmado por registro Hotmart"),
        ("Practitioner em PNL", "SBPNL", "Anterior a 2001", "—", "Formação confirmada pelo autor"),
        ("Practitioner em PNL", "Trivium", "Jun–dez/2001", "112h", "Formação documentada"),
        ("Master Practitioner em PNL", "Trivium", "Mar–mai/2001", "80h", "Formação documentada"),
    ]
    add_records_table(doc, ["Formação", "Instituição / professor", "Data / período", "Carga", "Natureza e uso"], rows, [1.48, 1.48, 1.12, 0.62, 1.72])

    doc.add_heading("7. Eixo 5 — Abordagens sistêmicas e constelações", level=1)
    rows = [
        ("Constelaciones Individuales en Consultorio", "USI", "11/09/2017", "120h", "Formação documentada"),
        ("Master Senior em Constelaciones Integrativas", "USI", "11/09/2017", "252h", "Título institucional"),
        ("Constelação Familiar", "USI", "30/06–02/07/2016", "4 dias", "Formação documentada"),
        ("Constelação em Consultório", "USI", "07–10/07/2016", "4 dias", "Formação documentada"),
        ("Constelação em Consultório III", "USI", "23–26/11/2016", "4 dias", "Formação documentada"),
        ("Constelação Organizacional — módulo I", "USI", "02–04/11/2016", "3 dias", "Formação documentada"),
        ("Constelações Integrativas VI e VII / Estruturais I", "USI", "19–25/04/2017", "42h", "Formação documentada"),
        ("Constelaciones Familiares I", "Elsever Institute / USI; Gabriel de Velasco", "12/02/2016", "56h", "Formação documentada"),
    ]
    add_records_table(doc, ["Formação", "Instituição / professor", "Data / período", "Carga", "Natureza e uso"], rows, [1.65, 1.55, 1.15, 0.65, 1.54])

    doc.add_heading("8. Eixo 6 — Consciência, espiritualidade, meditação e repertórios históricos", level=1)
    rows = [
        ("Atenção e concentração nas práticas meditativas — 3 módulos", "Associação Palas Athena", "Módulos 2 e 3 em 2013; sequência completa", "—", "Percurso integral confirmado"),
        ("Usui Shiki Ryoho — 1º grau", "Formação Reiki", "22/08/1993", "—", "Genealogia histórica"),
        ("Usui Shiki Ryoho — 2º grau", "Formação Reiki", "21/11/1993", "—", "Genealogia histórica"),
        ("Reiki Master", "Clínica Integrativa Michael / Grupo Júlia Magalhães", "30/09/2000", "—", "Genealogia histórica"),
        ("Registered Karuna Reiki Master", "International Center for Reiki Training", "03/06/2001", "—", "Registro BRM037"),
        ("Karuna Ki–Tera Mai — níveis I e II", "Formação Reiki", "Data parcial", "—", "Genealogia histórica"),
        ("Radiestesia — básico", "Curso livre", "13/09/1998", "15h", "Repertório histórico"),
        ("Radiestesia", "Antonio Rodrigues / equipe", "14/04/2002", "90h", "Repertório histórico"),
        ("Bioenergologia com Formação em Bioenergopatia", "Instituto Medeiros", "Conclusão ago/2004; declaração 29/11/2004", "114 unidades-aula", "Nota 9,0; sem conversão automática para horas"),
        ("Leader Training", "Núcleo Renascer", "23/10/1999", "—", "Desenvolvimento pessoal histórico"),
        ("Certificado de graduação em Aikido", "FEPAI", "13/06/1999", "—", "Graduação não publicada por legibilidade"),
        ("Certificado de graduação em Wado-Ryu Karatê-Dô", "Wado-Ryu Renmei do Brasil", "14/09/1982", "—", "Graduação não publicada por legibilidade"),
        ("Sistema de Equilíbrio Energético", "PHVIDA / Carlos Florêncio", "—", "—", "Formação como Healer"),
        ("Apometria Quântica — A Tecnologia do Espírito", "Carina Greco", "—", "—", "Repertório histórico"),
        ("Tameana — níveis 1 e 2 e Pojá", "Pleiadian Connection", "28/09/2018", "—", "Repertório histórico"),
        ("The Melchizedek Method — níveis 1 e 2", "Formação internacional", "12/10/2003", "—", "Practitioner Certificate"),
        ("Evidências Científicas sobre a Existência de Deus", "Instituto/Editora Aleph; Amit Goswami", "25/08/2007", "5h", "Palestra"),
        ("A Física da Alma", "Instituto/Editora Aleph; Amit Goswami", "31/05/2008", "3h", "Palestra"),
        ("Física Quântica, Múltiplas Dimensões e Universos Paralelos", "Instituto Aleph", "14–15/06/2013", "10h", "Workshop"),
        ("Psicologia Quântica Integrativa", "Instituto Aleph; Amit Goswami", "17/09/2013", "8h", "Workshop"),
        ("Recuperação de Memórias em Experiências Extraordinárias ou Anômalas", "Gilda Moura", "21–22/10/2017", "16h", "Formação em experiências incomuns"),
    ]
    add_records_table(doc, ["Formação", "Instituição / professor", "Data / período", "Carga", "Natureza e uso"], rows, [1.50, 1.48, 1.12, 0.62, 1.69])

    doc.add_heading("9. Eixo 7 — Docência, palestras, instituição e projetos autorais", level=1)
    rows = [
        ("Cursos de Reiki", "Cursos livres", "Período anterior à docência on-line", "Múltiplas turmas", "Ensino de práticas integrativas"),
        ("Curso de Seitai e Quiropraxia", "Shiozawa", "Período de atuação corporal", "—", "Experiência didática"),
        ("Corpo e Mente: um mesmo sistema", "PUC-SP, Faculdade de Psicologia", "01/10/2002", "Aula convidada", "Certificado institucional; sem vínculo docente regular"),
        ("Professor assistente de Gilda Moura", "Centralma", "2018/2019", "Turma presencial", "Hipnose, memória e experiências anômalas"),
        ("Cursos on-line de auto-hipnose", "Com Gilda Moura", "2020–2025", "3 cursos", "Docência digital"),
        ("Além das Fronteiras do Saber", "Com Gilda Moura", "A partir de 2020", "5 edições", "Professor e cocriador"),
        ("Workshop ABBSP", "Associação Brasileira de Brainspotting", "07/03/2026", "—", "Autorregulação, presença clínica e estados ampliados"),
        ("Palestra IGM", "Com Daiana Leite", "Ago/2026", "—", "Fenômenos anômalos e mudança da percepção da realidade"),
        ("CEEP", "Comunidade de Estudos em Espiritualidade e Psicoterapia", "2026–", "Eixo Consciência", "Professor; fenomenologia, discernimento e integração"),
        ("Traduzindo o Ser Humano", "Projeto autoral", "Em desenvolvimento", "—", "Criador, professor e autor"),
        ("Bioquantum / Centralma", "Empreendimento próprio", "Atuação presencial até 2020; clínica on-line atual", "—", "Fundador; atendimentos, cursos e organização institucional"),
    ]
    add_records_table(doc, ["Atividade", "Instituição / parceria", "Data / período", "Escala", "Natureza e uso"], rows, [1.48, 1.42, 1.28, 0.76, 1.69])

    doc.add_heading("10. Supervisão, prática e continuidade clínica", level=1)
    add_bullets(doc, [
        "O atendimento terapêutico começou nas formações corporais e integrativas e foi mantido de forma contínua.",
        "A prática psicoterapêutica começou em paralelo às formações em Psicanálise Integrativa, com estágios e atendimentos desde o início do percurso.",
        "A Psicologia consolidou a identidade profissional e o registro no CRP em 2010.",
        "A supervisão mensal com Mário Salvador permanece como formação clínica continuada em 2026.",
        "Supervisões, estágios e prática clínica são registrados separadamente de cargas horárias de cursos.",
    ])

    doc.add_heading("11. Regras de carga horária e não duplicação", level=1)
    add_bullets(doc, [
        "Não somar páginas duplicadas do mesmo certificado.",
        "Não somar níveis já incluídos em um programa geral.",
        "Não somar os certificados gerais do CEMETC aos módulos correspondentes sem análise de sobreposição.",
        "Não somar o Nível II Aleces de 90 horas ao programa total de 216 horas.",
        "Não somar registros de EMDR que possam ser versões diferentes do mesmo treinamento sem conferência documental.",
        "Não somar formações USI possivelmente incluídas em títulos institucionais mais amplos.",
        "Manter as 100 horas do Curso Intensivo de Shiatsu separadas das 2.000 horas de estágio no Shiozawa.",
        "Não converter as 114 unidades-aula do programa de Bioenergologia em 114 horas, pois a declaração não estabelece equivalência horária total.",
    ])

    doc.add_heading("12. Hierarquia pública recomendada", level=1)
    doc.add_heading("Núcleo frontal", level=2)
    add_bullets(doc, [
        "Psicólogo clínico e psicoterapeuta; CRP 06/98630.",
        "Mais de 30 anos no cuidado terapêutico, das práticas corporais e integrativas à consolidação com a Psicologia.",
        "Corpo, trauma, consciência e espiritualidade.",
        "Psicologia Transpessoal, Neurociências e Comportamento, Psicanálise Integrativa, EMDR, Brainspotting, hipnose e trauma.",
    ])
    doc.add_heading("Sustentação", level=2)
    add_bullets(doc, [
        "Supervisão clínica continuada; Neurobiologia Interpessoal e regulação autonômica.",
        "Experiência docente, TCC e produção intelectual em desenvolvimento.",
        "CEEP, ABBSP, IGM, Além das Fronteiras do Saber e Traduzindo o Ser Humano.",
        "Centralma e experiência de organização de cursos e projetos.",
    ])
    doc.add_heading("Linhagem histórica", level=2)
    add_bullets(doc, [
        "EOMA, Shiozawa, MTC, acupuntura, Qi Gong, Reiki, PNL, Seitai, Quiropraxia, constelações e repertórios correlatos.",
        "Essa camada explica a origem do olhar incorporado e interdisciplinar, sem ocupar a abertura de toda apresentação.",
    ])

    doc.add_heading("13. Formulações canônicas", level=1)
    add_key_value_table(doc, [
        ("Trajetória", "Mais de 30 anos de trajetória no cuidado terapêutico, iniciada nas práticas corporais e integrativas e consolidada com a Psicologia."),
        ("Transpessoal", "Pós-graduação em Psicologia Transpessoal pela FACIS — Faculdade de Ciências da Saúde de São Paulo. Formação concluída em 2010, com TCC entregue e certificado emitido em 2011."),
        ("Hipnose", "Hipnose clínica e terapêutica: formação continuada com estudos junto a Alberto Dell’Isola, Marco Ceda Natali, Fernando Toledo Piza e Sofia Bauer."),
        ("Docência", "Professor, palestrante e facilitador em cursos livres e educação continuada."),
        ("CEMETC–Yunnan", "Concluiu o Máster Iberoamericano en Acupuntura Bioenergética y Moxibustión, certificado em outubro de 2015 no âmbito institucional da Universidad de Medicina Tradicional China de Yunnan, do CEMETC e da Fundación Europea de Medicina Tradicional China."),
        ("Assinatura", "Corpo, trauma, consciência e espiritualidade na compreensão e tradução da experiência humana."),
    ])

    doc.add_heading("14. Governança para atualizações futuras", level=1)
    add_bullets(doc, [
        "A correção explícita mais recente de Fabiano prevalece sobre formulações anteriores produzidas por IA.",
        "O dossiê detalhado preserva todos os marcos; versões públicas selecionam apenas os eixos pertinentes ao contexto.",
        "Formação, supervisão, prática, docência, vínculo institucional e produção acadêmica permanecem em categorias distintas.",
        "Credenciais com validade administrativa encerrada permanecem como marcos históricos da trajetória; vigência atual é uma informação separada.",
        "Certificações confirmadas por Fabiano não devem ser reabertas como pendência de digitalização.",
        "O título e a carga de cada curso devem ser transcritos literalmente quando o uso externo exigir comprovação específica.",
    ])
    p = doc.add_paragraph(style="Small Note")
    p.add_run(
        "Fonte: vault Obsidian sincronizado no GitHub — Governança Curricular, Matriz Curricular Detalhada, Currículo Profissional Consolidado, Currículo Público Ampliado, Dossiê Curricular Mestre, auditoria principal de 117 páginas e complemento documental de 12 páginas (AUT-DOC-23). Branch de revisão: codex/atualiza-curriculo-brainspotting-palas-mtc-autohipnose-20260821."
    )

    path = OUT / "Dossie_Curricular_Mestre_Fabiano_Deliberalli_2026.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build_curriculum())
    print(build_dossier())
